import sys
from docx import Document

def extract_text_from_docx(input_docx, output_txt):
    # Word 파일 열기
    doc = Document(input_docx)
    
    # 출력할 텍스트 파일 열기
    with open(output_txt, 'w', encoding='utf-8') as f:
        # 문서의 모든 단락을 순회하면서 텍스트 추출
        for para in doc.paragraphs:
            f.write(para.text.strip() + "\n")
        
        # 문서에 있는 모든 표의 텍스트도 추출
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells]
                f.write("\t".join(row_text) + "\n")

if __name__ == "__main__":
    if len(sys.argv) != 3:
        print("사용법: python docx2txt.py <input_docx> <output_txt>")
    else:
        input_docx = sys.argv[1]
        output_txt = sys.argv[2]
        extract_text_from_docx(input_docx, output_txt)
        print(f"텍스트가 {output_txt} 파일에 저장되었습니다.")
