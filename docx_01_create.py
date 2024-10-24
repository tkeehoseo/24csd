from docx import Document
def create_example_docx(docx_path):
    doc = Document()
    doc.add_paragraph("첫 번째 단락입니다.")
    doc.add_paragraph("두 번째 단락입니다.")
    table = doc.add_table(rows=1, cols=2)
    row = table.rows[0]
    row.cells[0].text = "표 내용 1"
    row.cells[1].text = "표 내용 2"
    doc.save(docx_path)
docx_path = 'example.docx'
create_example_docx(docx_path)
print(f"{docx_path} 파일이 생성되었습니다.")
