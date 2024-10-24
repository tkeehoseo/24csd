from docx import Document
def extract_text_from_docx(docx_path):
    doc = Document(docx_path)
    doc_text = []
    for para in doc.paragraphs:
        doc_text.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                doc_text.append(cell.text)
    return '\n'.join(doc_text)
docx_path = 'example.docx'
extracted_text = extract_text_from_docx(docx_path)
print(extracted_text)
